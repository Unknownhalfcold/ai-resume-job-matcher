from __future__ import annotations

import io
from pathlib import Path


class DocumentParseError(RuntimeError):
    """Raised when an uploaded document cannot be parsed."""


def compact_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compacted: list[str] = []

    for line in lines:
        if line:
            compacted.append(line)
        elif compacted and compacted[-1]:
            compacted.append("")

    return "\n".join(compacted).strip()


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("python-docx is not installed.") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Unable to read DOCX file.") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return compact_text("\n".join(parts))


def extract_pdf_text(content: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise DocumentParseError("PyMuPDF is not installed.") from exc

    try:
        document = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise DocumentParseError("Unable to read PDF file.") from exc

    parts: list[str] = []
    for page in document:
        page_text = page.get_text("text")
        if page_text.strip():
            parts.append(page_text.strip())

    return compact_text("\n\n".join(parts))


def extract_resume_text(filename: str, content: bytes) -> tuple[str, list[str]]:
    extension = Path(filename).suffix.lower()
    warnings: list[str] = []

    if extension == ".docx":
        text = extract_docx_text(content)
    elif extension == ".pdf":
        text = extract_pdf_text(content)
        if not text:
            warnings.append("PDF 中没有提取到可复制文本，可能是扫描版 PDF，需要 OCR。")
    else:
        raise DocumentParseError("Only .docx and .pdf files are supported.")

    if not text:
        raise DocumentParseError("No text could be extracted from the uploaded file.")

    return text, warnings
